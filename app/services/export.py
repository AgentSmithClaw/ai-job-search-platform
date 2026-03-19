from __future__ import annotations

import io
from datetime import datetime
from typing import Optional
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib import colors


class ExportService:
    @staticmethod
    def generate_resume_docx(
        resume_draft: str,
        target_role: str,
        match_score: int,
        user_name: str = "",
    ) -> bytes:
        doc = Document()
        
        title = doc.add_heading(f"简历草稿 - {target_role}", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta.add_run(f"匹配度: {match_score}% | 生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}").font.size = Pt(10)
        
        doc.add_paragraph()
        
        for line in resume_draft.split('\n'):
            line = line.strip()
            if not line:
                doc.add_paragraph()
                continue
            
            if line.startswith('#') or line.startswith('【') and line.endswith('】'):
                heading_level = 2 if line.startswith('# ') else 2
                clean_line = line.lstrip('# ').strip()
                doc.add_heading(clean_line, level=heading_level)
            elif line.startswith('- '):
                doc.add_paragraph(clean_line if (clean_line := line[2:]) else '', style='List Bullet')
            elif line.startswith('• '):
                doc.add_paragraph(clean_line if (clean_line := line[2:]) else '', style='List Bullet')
            else:
                p = doc.add_paragraph()
                p.add_run(line)
        
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.add_run("—— 由 AI Job Search Platform 生成 ——").font.size = Pt(9)
        footer.runs[0].font.color.rgb = RGBColor(128, 128, 128)
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    @staticmethod
    def generate_analysis_report_pdf(
        target_role: str,
        match_score: int,
        summary: str,
        strengths: list[str],
        risks: list[str],
        gaps: list[dict],
        learning_plan: list[str],
        interview_focus: list[str],
        resume_suggestions: list[dict],
        resume_draft: str,
        user_name: str = "",
    ) -> bytes:
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=0.75 * inch,
            leftMargin=0.75 * inch,
            topMargin=0.75 * inch,
            bottomMargin=0.75 * inch,
        )
        
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            alignment=1,
            spaceAfter=20,
        )
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            spaceAfter=10,
            spaceBefore=15,
            textColor=colors.HexColor('#0a8a72'),
        )
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['Normal'],
            fontSize=10,
            spaceAfter=8,
            leading=14,
        )
        bullet_style = ParagraphStyle(
            'BulletStyle',
            parent=styles['Normal'],
            fontSize=10,
            leftIndent=20,
            spaceAfter=5,
            leading=13,
        )
        
        story = []
        
        story.append(Paragraph(f"求职差距分析报告 - {target_role}", title_style))
        story.append(Paragraph(f"匹配度: {match_score}%", body_style))
        story.append(Paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}", body_style))
        story.append(Spacer(1, 20))
        
        story.append(Paragraph("总体评估", heading_style))
        story.append(Paragraph(summary, body_style))
        
        story.append(Paragraph("优势分析", heading_style))
        for s in strengths:
            story.append(Paragraph(f"• {s}", bullet_style))
        
        story.append(Paragraph("风险提示", heading_style))
        for r in risks:
            story.append(Paragraph(f"• {r}", bullet_style))
        
        story.append(Paragraph("差距项分析", heading_style))
        for i, gap in enumerate(gaps, 1):
            story.append(Paragraph(f"<b>{i}. {gap.get('requirement', '')}</b>", body_style))
            story.append(Paragraph(f"   类别: {gap.get('category', '')} | 严重程度: {gap.get('severity', '')}", bullet_style))
            story.append(Paragraph(f"   证据: {gap.get('evidence', '')}", bullet_style))
            story.append(Paragraph(f"   建议: {gap.get('recommendation', '')}", bullet_style))
            story.append(Spacer(1, 5))
        
        story.append(Paragraph("学习计划", heading_style))
        for i, plan in enumerate(learning_plan, 1):
            story.append(Paragraph(f"{i}. {plan}", bullet_style))
        
        story.append(Paragraph("面试准备重点", heading_style))
        for focus in interview_focus:
            story.append(Paragraph(f"• {focus}", bullet_style))
        
        story.append(Paragraph("简历优化建议", heading_style))
        for i, sug in enumerate(resume_suggestions, 1):
            story.append(Paragraph(f"<b>建议 {i}:</b> {sug.get('reason', '')}", body_style))
            story.append(Paragraph(f"原文: {sug.get('original', '')}", bullet_style))
            story.append(Paragraph(f"优化: {sug.get('optimized', '')}", bullet_style))
            story.append(Spacer(1, 5))
        
        story.append(PageBreak())
        story.append(Paragraph("岗位定制简历草稿", heading_style))
        for line in resume_draft.split('\n'):
            if line.strip():
                story.append(Paragraph(line, body_style))
            else:
                story.append(Spacer(1, 8))
        
        story.append(Spacer(1, 30))
        footer_text = "—— 由 AI Job Search Platform 生成 ——"
        footer_style = ParagraphStyle(
            'Footer',
            parent=styles['Normal'],
            fontSize=9,
            alignment=1,
            textColor=colors.grey,
        )
        story.append(Paragraph(footer_text, footer_style))
        
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()


export_service = ExportService()